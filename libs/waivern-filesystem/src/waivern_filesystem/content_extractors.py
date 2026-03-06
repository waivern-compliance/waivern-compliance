"""Content extractors for rich document formats (DOCX, XLSX).

Suffix-based dispatch: each extractor is a function (Path) -> str.
The CONTENT_EXTRACTORS dict maps file suffixes to extractor functions,
used by _read_file_content() in the connector to handle binary formats.

Dependencies are optional — guarded imports at module level satisfy PLC0415
while allowing the package to be used without python-docx or openpyxl.
"""

from collections.abc import Callable
from pathlib import Path

from waivern_core.errors import ConnectorExtractionError

_has_docx = False
try:
    from docx import Document

    _has_docx = True
except ImportError:
    pass

_has_openpyxl = False
try:
    from openpyxl import load_workbook

    _has_openpyxl = True
except ImportError:
    pass


def extract_docx(file_path: Path) -> str:
    """Extract text from a DOCX file.

    Extracts paragraphs and table cell text, concatenated with newlines.
    Headers and footers are excluded — they typically contain page numbers
    and classification markings that add noise without informational value.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Extracted text content.

    Raises:
        ConnectorExtractionError: If python-docx is not installed or the
            file cannot be read.

    """
    if not _has_docx:
        raise ConnectorExtractionError(
            "python-docx is required for .docx files. "
            "Install with: uv pip install waivern-filesystem[docs]"
        )

    try:
        doc = Document(str(file_path))  # pyright: ignore[reportPossiblyUnboundVariable]
    except Exception as e:
        raise ConnectorExtractionError(
            f"Failed to read DOCX file {file_path}: {e}"
        ) from e

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append("\t".join(cells))

    return "\n".join(parts)


def extract_xlsx(file_path: Path) -> str:
    """Extract text from an XLSX file as Markdown tables.

    Each sheet is rendered as a ``## Sheet: <name>`` heading followed by a
    Markdown table. The first data row becomes the header row, followed by
    a separator row (``| --- | --- |``), then the remaining data rows.
    Empty rows are skipped.

    Args:
        file_path: Path to the .xlsx file.

    Returns:
        Extracted text content in Markdown table format.

    Raises:
        ConnectorExtractionError: If openpyxl is not installed or the
            file cannot be read.

    """
    if not _has_openpyxl:
        raise ConnectorExtractionError(
            "openpyxl is required for .xlsx files. "
            "Install with: uv pip install waivern-filesystem[docs]"
        )

    try:
        wb = load_workbook(str(file_path), read_only=True, data_only=True)  # pyright: ignore[reportPossiblyUnboundVariable]
    except Exception as e:
        raise ConnectorExtractionError(
            f"Failed to read XLSX file {file_path}: {e}"
        ) from e

    sheet_parts: list[str] = []

    for sheet in wb.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows():
            cell_values = [
                str(cell.value) if cell.value is not None else "" for cell in row
            ]
            if any(v != "" for v in cell_values):
                rows.append(cell_values)

        if not rows:
            continue

        lines: list[str] = [f"## Sheet: {sheet.title}", ""]

        # First row as header
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")

        # Remaining rows as data
        for data_row in rows[1:]:
            lines.append("| " + " | ".join(data_row) + " |")

        sheet_parts.append("\n".join(lines))

    wb.close()
    return "\n\n".join(sheet_parts)


CONTENT_EXTRACTORS: dict[str, Callable[[Path], str]] = {
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
}

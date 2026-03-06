"""Tests for content extractors (DOCX and XLSX text extraction)."""

from pathlib import Path
from typing import Any

import pytest
from docx import Document
from openpyxl import Workbook
from waivern_core.errors import ConnectorExtractionError

from waivern_filesystem.content_extractors import extract_docx, extract_xlsx


def _create_workbook(
    tmp_path: Path, filename: str, sheets: dict[str, list[list[Any]]]
) -> Path:
    """Create an XLSX workbook with the given sheet data.

    Args:
        tmp_path: Temporary directory path.
        filename: Name of the file to create.
        sheets: Dict of sheet_name -> list of rows (each row is a list of values).

    Returns:
        Path to the created .xlsx file.

    """
    wb = Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            ws = wb.active
            assert ws is not None
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        for row in rows:
            ws.append(row)
    file_path = tmp_path / filename
    wb.save(str(file_path))
    return file_path


class TestExtractDocx:
    """Tests for DOCX text extraction."""

    # =========================================================================
    # Content extraction
    # =========================================================================

    def test_extract_docx_paragraphs(self, tmp_path: Path):
        """Paragraphs are extracted as newline-separated text."""
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("Third paragraph")
        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        result = extract_docx(file_path)

        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result

    def test_extract_docx_tables(self, tmp_path: Path):
        """Table cell text appears in extracted output."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Value 1"
        table.cell(1, 1).text = "Value 2"
        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        result = extract_docx(file_path)

        assert "Header 1" in result
        assert "Header 2" in result
        assert "Value 1" in result
        assert "Value 2" in result

    def test_extract_docx_paragraphs_and_tables_combined(self, tmp_path: Path):
        """Both paragraphs and tables appear when the document contains both."""
        doc = Document()
        doc.add_paragraph("Introduction paragraph")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Control"
        table.cell(0, 1).text = "Status"
        doc.add_paragraph("Conclusion paragraph")
        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        result = extract_docx(file_path)

        assert "Introduction paragraph" in result
        assert "Control" in result
        assert "Status" in result
        assert "Conclusion paragraph" in result

    # =========================================================================
    # Edge cases
    # =========================================================================

    def test_extract_docx_empty_document(self, tmp_path: Path):
        """Empty DOCX returns empty string, not an error."""
        doc = Document()
        file_path = tmp_path / "empty.docx"
        doc.save(str(file_path))

        result = extract_docx(file_path)

        assert result == ""

    def test_extract_docx_missing_package(self, monkeypatch):
        """ConnectorExtractionError with install message when python-docx is missing."""
        import waivern_filesystem.content_extractors as mod

        monkeypatch.setattr(mod, "_has_docx", False)

        with pytest.raises(ConnectorExtractionError, match="python-docx is required"):
            extract_docx(Path("irrelevant.docx"))

    def test_extract_docx_corrupt_file(self, tmp_path: Path):
        """Corrupt DOCX raises ConnectorExtractionError, not an unhandled exception."""
        file_path = tmp_path / "corrupt.docx"
        file_path.write_bytes(b"\x00\x01\x02\x03\xff\xfe\xfd")

        with pytest.raises(ConnectorExtractionError, match="Failed to read DOCX"):
            extract_docx(file_path)


class TestExtractXlsx:
    """Tests for XLSX text extraction."""

    # =========================================================================
    # Content extraction
    # =========================================================================

    def test_extract_xlsx_single_sheet(self, tmp_path: Path):
        """Single sheet produces heading + Markdown table with header and separator."""
        file_path = _create_workbook(
            tmp_path,
            "test.xlsx",
            {
                "Risk Register": [
                    ["ID", "Risk", "Status"],
                    ["R-001", "Data breach", "Open"],
                    ["R-002", "System failure", "Mitigating"],
                ],
            },
        )

        result = extract_xlsx(file_path)

        assert "## Sheet: Risk Register" in result
        assert "| ID | Risk | Status |" in result
        assert "| --- | --- | --- |" in result
        assert "| R-001 | Data breach | Open |" in result
        assert "| R-002 | System failure | Mitigating |" in result

    def test_extract_xlsx_multi_sheet(self, tmp_path: Path):
        """Both sheet names appear as headings in multi-sheet workbooks."""
        file_path = _create_workbook(
            tmp_path,
            "test.xlsx",
            {
                "Controls": [["Ref", "Name"], ["A.5.1", "Policies"]],
                "Assets": [["Asset", "Owner"], ["Database", "IT Ops"]],
            },
        )

        result = extract_xlsx(file_path)

        assert "## Sheet: Controls" in result
        assert "## Sheet: Assets" in result
        assert "| A.5.1 | Policies |" in result
        assert "| Database | IT Ops |" in result

    def test_extract_xlsx_empty_rows_skipped(self, tmp_path: Path):
        """Empty rows are not present in the output."""
        file_path = _create_workbook(
            tmp_path,
            "test.xlsx",
            {
                "Data": [["Header"], ["Row 1"], [None], ["Row 3"]],
            },
        )

        result = extract_xlsx(file_path)

        assert "| Row 1 |" in result
        assert "| Row 3 |" in result
        # Empty row should not produce a pipe-only line
        lines = result.strip().split("\n")
        for line in lines:
            if line.startswith("|"):
                stripped = line.replace("|", "").replace("-", "").strip()
                assert stripped != "" or "---" in line

    # =========================================================================
    # Edge cases
    # =========================================================================

    def test_extract_xlsx_empty_workbook(self, tmp_path: Path):
        """Empty XLSX (no data rows) returns empty string, not an error."""
        file_path = _create_workbook(tmp_path, "empty.xlsx", {"Empty": []})

        result = extract_xlsx(file_path)

        assert result == ""

    def test_extract_xlsx_missing_package(self, monkeypatch):
        """ConnectorExtractionError with install message when openpyxl is missing."""
        import waivern_filesystem.content_extractors as mod

        monkeypatch.setattr(mod, "_has_openpyxl", False)

        with pytest.raises(ConnectorExtractionError, match="openpyxl is required"):
            extract_xlsx(Path("irrelevant.xlsx"))

    def test_extract_xlsx_corrupt_file(self, tmp_path: Path):
        """Corrupt XLSX raises ConnectorExtractionError, not an unhandled exception."""
        file_path = tmp_path / "corrupt.xlsx"
        file_path.write_bytes(b"\x00\x01\x02\x03\xff\xfe\xfd")

        with pytest.raises(ConnectorExtractionError, match="Failed to read XLSX"):
            extract_xlsx(file_path)
